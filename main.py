'''Scraping product detail from www.drstaranandram.com website and convert that page into document to refrence purpose

  .......................Browser automation done by Selenoum Python Library .........................//

..............scraping the text ,title and image is done by bs4 liberary ........................//

..............dowladed copy image store in local file and in each pagination image is rewrite ................//

..............project is Done Adithyan M (Be Computer Science Engineering ) Coimbatore ,TamilNadu,India................//
.................................................All copy Rights............................'''

# import the liberay Fuction
import re
import time
from webdriver_manager.chrome import ChromeDriverManager
from selenium import webdriver
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
import requests
from selenium.webdriver.support import expected_conditions
from docx import Document

# download Chrome Drive
driver = webdriver.Chrome(ChromeDriverManager().install())
document = Document()
font_style = document.styles

'''This function download the image and store in the local directoey were the py file is situated and each pagination it rer write by new file 

.............requist the url in and write tha data in file '''


def img_download(url, i):
    try:
        img_data = requests.get(url).content
        with open(f'img{i}.jpg', 'wb') as handler:
            handler.write(img_data)
    except expected_conditions:
        print(expected_conditions)


# pagination will ocur in 14 times which is my maximum page in my website
for x in range(1, 15):

    # moving into page website
    try:
        driver.get(f'https://www.drstaranandram.com/shop/page/{x}/')
        driver.implicitly_wait(5000)
    # to reduice stale error must include with try bliock
    except expected_conditions:
        print("From drive", expected_conditions)

    '''selenium is working by grabing current loadded information
    
     or script in file to make all file loaded properly or to reduce
     
      stament error is better to put delay  iin the automation or before
      
       clincking any key '''
    time.sleep(2)

    # values are the clickable icon which is identified by class name
    values = (driver.find_elements(By.CLASS_NAME, 'shop-item-title-link'))

    # parsing content to bs4
    page = requests.get(driver.current_url)
    soup = BeautifulSoup(page.content, 'html.parser')
    link_url = []

    '''..............geting all images link from card division through 
    
    find_all method in bs4  and downloading each image through function 
   
    img_download below ...............'''

    for link in soup.find_all('img',
                              attrs={'src': re.compile("https://")}):
        # display the actual urls
        link_url.append(link.get('src'))
    for l in range(1, 13):
        img_download(link_url[l], l)

    for i in range(0, len(values)):
        values[i].click()
        time.sleep(1)
        page = requests.get(driver.current_url)
        soup = BeautifulSoup(page.content, 'html.parser')

        # extarcting the title of product
        try:
            try:
                title = soup.find('div',
                                  class_='summary entry-summary').text
            except AttributeError:
                title = ''
                print(AttributeError)

            # extracting the descrioption of product
            try:
                description = soup.find('div',
                                        class_='woocommerce-Tabs-panel woocommerce-Tabs-panel--description panel '
                                               'entry-content wc-tab').text
            except AttributeError:
                description = ''
                print(AttributeError)
            document.add_picture(f'img{i + 1}.jpg')

            # adding heading to document file

            document.add_heading(title, 3)

            # adding description to the document file

            paragraph = document.add_paragraph(description)
            document.add_page_break()
            document.save(f'cust{2}.docx')
        except expected_conditions:
            print("error ", expected_conditions)

        driver.back()
        time.sleep(1)
